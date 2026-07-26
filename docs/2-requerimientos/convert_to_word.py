# -*- coding: utf-8 -*-
"""
Conversor Markdown -> Word (.docx) para los documentos del TP.
Maneja: encabezados, tablas, listas (viñeta/numeradas), bloques de código
(los diagramas Mermaid se marcan para renderizar a imagen aparte),
blockquotes y formato inline (negrita, cursiva, código).
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_inline(paragraph, text):
    """Agrega runs interpretando **negrita**, *cursiva* y `código`."""
    pattern = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = 'Consolas'; r.font.size = Pt(9)
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_row(r):
    return [c.strip() for c in r.strip().strip('|').split('|')]


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ---- bloque de código ----
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i])
                i += 1
            i += 1  # cierre ```
            if lang == 'mermaid':
                p = doc.add_paragraph()
                rr = p.add_run('[Diagrama — renderizar a imagen desde el código siguiente '
                               '(mermaid.live / draw.io) e insertarlo aquí]')
                rr.italic = True
                rr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            for cl in code:
                p = doc.add_paragraph()
                rr = p.add_run(cl if cl else ' ')
                rr.font.name = 'Consolas'; rr.font.size = Pt(8)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            doc.add_paragraph()
            continue

        # ---- tabla ----
        if (stripped.startswith('|') and i + 1 < n and
                re.match(r'^\|?[\s:|\-]+\|?\s*$', lines[i + 1].strip()) and
                '-' in lines[i + 1]):
            header = parse_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(parse_row(lines[i].strip()))
                i += 1
            ncols = len(header)
            header_empty = all(h == '' for h in header)

            t = doc.add_table(rows=0, cols=ncols)
            t.style = 'Table Grid'

            if not header_empty:
                cells = t.add_row().cells
                for j in range(ncols):
                    cells[j].paragraphs[0].text = ''
                    add_inline(cells[j].paragraphs[0], header[j] if j < len(header) else '')
                    for run in cells[j].paragraphs[0].runs:
                        run.bold = True
                    shade_cell(cells[j], 'D9E2F3')
            for row in rows:
                cells = t.add_row().cells
                for j in range(ncols):
                    cells[j].paragraphs[0].text = ''
                    add_inline(cells[j].paragraphs[0], row[j] if j < len(row) else '')
            doc.add_paragraph()
            continue

        # ---- encabezados ----
        mh = re.match(r'^(#+)\s+(.*)$', stripped)
        if mh:
            level = len(mh.group(1))
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', mh.group(2))
            text = re.sub(r'`([^`]+)`', r'\1', text)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # ---- regla horizontal ----
        if stripped == '---':
            i += 1
            continue

        # ---- blockquote ----
        if stripped.startswith('>'):
            qlines = []
            while i < n and lines[i].strip().startswith('>'):
                qlines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            rr = p.add_run(' '.join(q for q in qlines if q))
            rr.italic = True
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # ---- viñeta ----
        mb = re.match(r'^\s*[-*]\s+(.*)$', line)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mb.group(1))
            i += 1
            continue

        # ---- numerada ----
        mn = re.match(r'^\s*\d+\.\s+(.*)$', line)
        if mn:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mn.group(1))
            i += 1
            continue

        # ---- vacío ----
        if not stripped:
            i += 1
            continue

        # ---- párrafo normal ----
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(base, 'word')
    os.makedirs(out, exist_ok=True)
    docs = [
        '01-requerimientos',
        '02-alcance',
        '03-diagramas-flujo-datos',
        '04-diagrama-entidad-relacion',
        '05-caso-de-uso-plantilla',
        '06-caso-de-uso-asistencia-automatica',
    ]
    for d in docs:
        md = os.path.join(base, d + '.md')
        dx = os.path.join(out, d + '.docx')
        convert(md, dx)
        print('OK ->', os.path.basename(dx))
